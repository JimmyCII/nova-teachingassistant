"""Weekly Quiz Generator Tool."""
import os
from pathlib import Path
from google import genai
import docx
from agent.tools.pedagogy_critic import review_and_revise

_MODEL = os.getenv("HOMEWORK_MODEL", "gemini-2.5-flash")

_PROMPT = """You create a Weekly Quiz for a 6th-grade Arizona math class based on the provided standards.
Format the output as clean text/Markdown.
Include EXACTLY 4-5 questions targeting DOK 2 level (Skill/Concept: apply, graph, classify, estimate, compare).
Provide an Answer Key at the very bottom.
Do NOT use any real student names. Use fictional names if needed.

Standards to cover: {standards}
Quiz Title: {title}
"""

def generate_weekly_quiz(standards: list[str], title: str) -> dict:
    """Generate a weekly quiz using Gemini, critique it, and save as a Word document."""
    # 1. Generator Agent
    prompt = _PROMPT.format(standards=", ".join(standards), title=title)
    
    client = genai.Client(api_key=os.environ.get("GOOGLE_API_KEY", ""))
    response = client.models.generate_content(model=_MODEL, contents=prompt)
    draft_text = response.text
    
    # 2. Pedagogy Critic Agent (Multi-Agent Pipeline)
    print("Sending draft to Pedagogy Critic...")
    result = review_and_revise(draft_text, content_type="quiz", target_dok=2)
    
    print(f"Critic Verdict: {result.get('verdict', 'unknown')}")
    if result.get("issues"):
        print(f"Critic Issues Addressed: {result['issues']}")
    print(f"Standards Source: {result.get('standards_source', 'unknown')}")
        
    revised_text = result.get("revised_text", draft_text)
    
    # 3. Output Formatter (Word .docx)
    out_dir = Path("docs/02_Generated/Quizzes")
    out_dir.mkdir(parents=True, exist_ok=True)
    
    safe_title = title.replace("/", "-").replace("\\", "-").replace(" ", "_")
    out_path = out_dir / f"{safe_title}.docx"
    
    doc = docx.Document()
    doc.add_heading(title, 0)
    
    for line in revised_text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == '':
            continue
        else:
            doc.add_paragraph(line)
            
    doc.save(out_path)
    print(f"Saved approved quiz to {out_path}")
        
    # 4. Save to Drive Backbone & Comms Loop
    try:
        from agent.tools.spiral_homework.drive_store import GoogleDriveClient
        from agent.tools.request_logger import log_nova_task
        from mcp_servers.comms_server.server import send_draft_for_approval
        
        drive_client = GoogleDriveClient()
        root = drive_client.ensure_folder("Nova Teaching Assistant", None)
        quizzes_folder = drive_client.ensure_folder("Quizzes", root)
        drive_client.ensure_folder("02_Approved", quizzes_folder) # Create empty approved folder so it exists
        drafts_folder = drive_client.ensure_folder("01_Drafts", quizzes_folder)
        
        link, file_id = drive_client.upload_file(str(out_path), drafts_folder, out_path.name)
        
        karrie_email = os.environ.get("KARRIE_EMAIL", "test@example.com")
        send_res = send_draft_for_approval(karrie_email, link, out_path.name, title)
        
        log_res = log_nova_task(title, ", ".join(standards), status="Pending Approval", file_id=file_id)
        
        return {
            "status": "success",
            "local_path": str(out_path),
            "drive_link": link,
            "message": f"Quiz generated & uploaded: {link}. {send_res}. {log_res}"
        }
    except Exception as e:
        return {
            "status": "success",
            "local_path": str(out_path),
            "message": f"Quiz generated locally to {out_path}, but Drive upload failed: {e}"
        }
